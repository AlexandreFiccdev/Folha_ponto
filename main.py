import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
import os
import re
import numpy as np

# Carregar a planilha
caminho_planilha = "C:/Users/xande/Documents/projeto_mae/Livro_março.xlsx"
df = pd.read_excel(caminho_planilha, engine="openpyxl")

# Converter os nomes das colunas para strings
df.columns = [str(col) if not isinstance(col, int) else col for col in df.columns]

# Carregar o modelo de documento
caminho_modelo = "C:/Users/xande/Documents/projeto_mae/MARÇO.docx"
modelo_doc = Document(caminho_modelo)

# Criar diretório para salvar os arquivos
pasta_matutino = "C:/Users/xande/Documents/projeto_mae/Livro_ponto_matutino"
pasta_vespertino = "C:/Users/xande/Documents/projeto_mae/Livro_ponto_vespertino"
os.makedirs(pasta_matutino, exist_ok=True)
os.makedirs(pasta_vespertino, exist_ok=True)

# Mapear colunas do Excel para os campos do documento
colunas = {
    "NOME": "NOME",
    "CADASTRO": "CADASTRO",
    "VINCULO": "VINCULO",
    "MATÉRIA": "MATÉRIA",
    "TURNO": "TURNO",
    "SITUAÇÃO": "SITUAÇÃO",
}

# Mapeamento de colunas das aulas por dia da semana
mapeamento_aulas = {
    "SEG": [21, 22, 23, 24],
    "TER": [31, 32, 33, 34],
    "QUA": [41, 42, 43, 44],
    "QUI": [51, 52, 53, 54],
    "SEX": [61, 62, 63, 64],
}

# Lista de dias úteis, ignorando sábados (SAB) e domingos (DOM)
dias_uteis = []
dia_semana = 0 

for dia in range(3, 29):  # Começar do dia 3, pois dia 1 e 2 são sábado e domingo
    if dia_semana < 5: 
        dias_uteis.append((dia, list(mapeamento_aulas.keys())[dia_semana]))
    dia_semana = (dia_semana + 1) % 7  # Avançar um dia na semana

def encontrar_tabela_aulas(doc):
    """ Encontra a tabela correta no documento verificando os marcadores numéricos. """
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto = celula.text.strip()
                # Tenta encontrar números ignorando formatação como "1º"
                if re.match(r'\d+', texto):
                    numero_dia = int(re.search(r'\d+', texto).group())  # Pega apenas os números
                    if numero_dia == 3:  # Começa a partir do dia 3
                        return tabela
    return None


def limpar_e_substituir_texto(celula, texto):
    """ Limpa a célula e adiciona o novo texto corretamente formatado. """
    celula.text = ""
    p = celula.paragraphs[0]
    p.clear()
    run = p.add_run(str(texto))
    run.font.size = Pt(8) 
    p.alignment = 1  

def preencher_documento(professor):
    doc = Document(caminho_modelo)

    # Substituir valores no cabeçalho (parágrafos e células de tabelas)
    for paragrafo in doc.paragraphs:
        for chave, coluna in colunas.items():
            if f"«{chave}»" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(f"«{chave}»", str(professor.get(coluna, '--')))

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, coluna in colunas.items():
                    if f"«{chave}»" in celula.text:
                        limpar_e_substituir_texto(celula, professor.get(coluna, '--'))

    # Identificar a tabela correta
    tabela_aulas = encontrar_tabela_aulas(doc)
    if not tabela_aulas:
        print(f"Erro: Tabela de horários não encontrada para {professor.get('NOME', 'Desconhecido')}")
        return None

    # Preencher horários das aulas na tabela corretamente
    for linha in tabela_aulas.rows:
        celulas = linha.cells
        primeira_celula = celulas[0].text.strip()

        if primeira_celula.isdigit():  
            dia = int(primeira_celula)

            if dia in [d[0] for d in dias_uteis]:  
                dia_semana = [d[1] for d in dias_uteis if d[0] == dia][0]  
                colunas_aulas = mapeamento_aulas[dia_semana]

                for i, coluna in enumerate(colunas_aulas):
                    valor = professor.get(coluna, '--')
                    if pd.isna(valor):  
                        valor = "--"
                    limpar_e_substituir_texto(celulas[i + 1], valor)



    # Criar nome de arquivo válido
    nome_professor = str(professor.get('NOME', 'Sem_Nome'))
    nome_professor = re.sub(r'[^a-zA-Z0-9_]', '_', nome_professor)
    nome_arquivo = os.path.join(pasta_saida, f"Folha_{nome_professor}.docx")

    # Salvar arquivo individual
    doc.save(nome_arquivo)
    return nome_arquivo

arquivos_gerados = []
for _, professor in df.iterrows():
    nome_professor = professor.get('NOME', 'Sem_Nome')
    turno = professor.get('TURNO', 'Desconhecido')  
    
    # Determinar a pasta correta com base no turno
    if "MATUTINO" in turno.upper():
        pasta_saida = pasta_matutino
    elif "VESPERTINO" in turno.upper():
        pasta_saida = pasta_vespertino
    else:
        pasta_saida = "C:/Users/xande/Documents/projeto_mae/Livro_ponto_outros" 
    
    os.makedirs(pasta_saida, exist_ok=True) 
    
    nome_professor_limpo = re.sub(r'[^a-zA-Z0-9_]', '_', nome_professor)
    caminho_arquivo = os.path.join(pasta_saida, f"Folha_{nome_professor_limpo}.docx")
    
    caminho_arquivo = preencher_documento(professor)
    if caminho_arquivo:
        arquivos_gerados.append(caminho_arquivo)

print("Arquivos gerados:", arquivos_gerados)    

